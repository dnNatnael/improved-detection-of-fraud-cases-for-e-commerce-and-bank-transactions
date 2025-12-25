"""
Generate Comprehensive Project Report in DOCX Format

This script creates a detailed report of the fraud detection project,
including visualizations, charts, tables, and analysis results.
"""

import sys
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import warnings
warnings.filterwarnings('ignore')

# Import SMOTE and train_test_split at module level to avoid import issues
try:
    from sklearn.model_selection import train_test_split
    from imblearn.over_sampling import SMOTE
except ImportError as e:
    print(f"Warning: Could not import SMOTE or train_test_split: {e}")
    train_test_split = None
    SMOTE = None

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(__file__)), 'src'))

# Import project modules
from data_loading import load_data
from data_cleaning import clean_data, standardize_datatypes
from feature_engineering import engineer_features, merge_geolocation
from eda_utils import analyze_class_distribution, get_data_summary

# Set style
plt.style.use('seaborn-v0_8-darkgrid')
sns.set_palette("husl")


def create_visualizations(output_dir='temp_charts'):
    """Create all visualizations for the report"""
    os.makedirs(output_dir, exist_ok=True)
    
    # Base paths
    base_path = os.path.dirname(os.path.dirname(__file__))
    data_raw = os.path.join(base_path, 'data', 'raw')
    data_processed = os.path.join(base_path, 'data', 'processed')
    
    chart_paths = {}
    
    print("📊 Loading datasets...")
    # Load raw data for EDA
    try:
        fraud_raw = load_data(os.path.join(data_raw, 'Fraud_Data.csv'))
        creditcard_raw = load_data(os.path.join(data_raw, 'creditcard.csv'))
        ip_df = load_data(os.path.join(data_raw, 'IpAddress_to_Country.csv'))
        
        # Load processed data if available
        fraud_processed = None
        if os.path.exists(os.path.join(data_processed, 'fraud_data_cleaned.csv')):
            fraud_processed = load_data(os.path.join(data_processed, 'fraud_data_cleaned.csv'))
    except Exception as e:
        print(f"⚠ Error loading data: {e}")
        return chart_paths
    
    # 1. Class Distribution - Fraud Data
    print("📈 Creating class distribution charts...")
    if 'class' in fraud_raw.columns:
        fig, axes = plt.subplots(1, 2, figsize=(12, 5))
        
        class_dist = fraud_raw['class'].value_counts().sort_index()
        class_dist.plot(kind='bar', ax=axes[0], color=['#2ecc71', '#e74c3c'])
        axes[0].set_title('Fraud Data - Class Distribution', fontsize=14, fontweight='bold')
        axes[0].set_xlabel('Class (0=Legitimate, 1=Fraud)', fontsize=12)
        axes[0].set_ylabel('Count', fontsize=12)
        axes[0].tick_params(axis='x', rotation=0)
        axes[0].grid(axis='y', alpha=0.3)
        
        class_dist.plot(kind='pie', ax=axes[1], autopct='%1.2f%%', 
                       colors=['#2ecc71', '#e74c3c'], startangle=90)
        axes[1].set_title('Fraud Data - Class Distribution (%)', fontsize=14, fontweight='bold')
        axes[1].set_ylabel('')
        
        plt.tight_layout()
        chart_paths['fraud_class_dist'] = os.path.join(output_dir, 'fraud_class_dist.png')
        plt.savefig(chart_paths['fraud_class_dist'], dpi=300, bbox_inches='tight')
        plt.close()
    
    # 2. Class Distribution - Credit Card
    if 'Class' in creditcard_raw.columns:
        fig, axes = plt.subplots(1, 2, figsize=(12, 5))
        
        class_dist_cc = creditcard_raw['Class'].value_counts().sort_index()
        class_dist_cc.plot(kind='bar', ax=axes[0], color=['#3498db', '#e74c3c'])
        axes[0].set_title('Credit Card Data - Class Distribution', fontsize=14, fontweight='bold')
        axes[0].set_xlabel('Class (0=Legitimate, 1=Fraud)', fontsize=12)
        axes[0].set_ylabel('Count', fontsize=12)
        axes[0].tick_params(axis='x', rotation=0)
        axes[0].grid(axis='y', alpha=0.3)
        
        class_dist_cc.plot(kind='pie', ax=axes[1], autopct='%1.2f%%', 
                          colors=['#3498db', '#e74c3c'], startangle=90)
        axes[1].set_title('Credit Card Data - Class Distribution (%)', fontsize=14, fontweight='bold')
        axes[1].set_ylabel('')
        
        plt.tight_layout()
        chart_paths['creditcard_class_dist'] = os.path.join(output_dir, 'creditcard_class_dist.png')
        plt.savefig(chart_paths['creditcard_class_dist'], dpi=300, bbox_inches='tight')
        plt.close()
    
    # 3. Purchase Value Distribution
    print("📈 Creating purchase value distribution...")
    if 'purchase_value' in fraud_raw.columns:
        fig, axes = plt.subplots(1, 2, figsize=(14, 5))
        
        fraud_raw['purchase_value'].hist(bins=50, ax=axes[0], edgecolor='black', color='#3498db')
        axes[0].set_title('Purchase Value Distribution (Fraud Data)', fontsize=14, fontweight='bold')
        axes[0].set_xlabel('Purchase Value', fontsize=12)
        axes[0].set_ylabel('Frequency', fontsize=12)
        axes[0].grid(axis='y', alpha=0.3)
        
        # By class
        fraud_legit = fraud_raw[fraud_raw['class'] == 0]['purchase_value']
        fraud_fraud = fraud_raw[fraud_raw['class'] == 1]['purchase_value']
        
        axes[1].hist([fraud_legit, fraud_fraud], bins=50, label=['Legitimate', 'Fraud'], 
                    color=['#2ecc71', '#e74c3c'], alpha=0.7, edgecolor='black')
        axes[1].set_title('Purchase Value by Class', fontsize=14, fontweight='bold')
        axes[1].set_xlabel('Purchase Value', fontsize=12)
        axes[1].set_ylabel('Frequency', fontsize=12)
        axes[1].legend()
        axes[1].grid(axis='y', alpha=0.3)
        
        plt.tight_layout()
        chart_paths['purchase_value_dist'] = os.path.join(output_dir, 'purchase_value_dist.png')
        plt.savefig(chart_paths['purchase_value_dist'], dpi=300, bbox_inches='tight')
        plt.close()
    
    # 4. Age Distribution
    if 'age' in fraud_raw.columns:
        fig, ax = plt.subplots(figsize=(10, 6))
        fraud_raw['age'].hist(bins=30, ax=ax, edgecolor='black', color='#f39c12')
        ax.set_title('Age Distribution (Fraud Data)', fontsize=14, fontweight='bold')
        ax.set_xlabel('Age', fontsize=12)
        ax.set_ylabel('Frequency', fontsize=12)
        ax.grid(axis='y', alpha=0.3)
        
        plt.tight_layout()
        chart_paths['age_dist'] = os.path.join(output_dir, 'age_dist.png')
        plt.savefig(chart_paths['age_dist'], dpi=300, bbox_inches='tight')
        plt.close()
    
    # 5. Credit Card Amount Distribution
    if 'Amount' in creditcard_raw.columns:
        fig, axes = plt.subplots(1, 2, figsize=(14, 5))
        
        creditcard_raw['Amount'].hist(bins=50, ax=axes[0], edgecolor='black', color='#9b59b6')
        axes[0].set_title('Transaction Amount Distribution', fontsize=14, fontweight='bold')
        axes[0].set_xlabel('Amount', fontsize=12)
        axes[0].set_ylabel('Frequency', fontsize=12)
        axes[0].set_yscale('log')
        axes[0].grid(axis='y', alpha=0.3)
        
        # Log transformed
        np.log1p(creditcard_raw['Amount']).hist(bins=50, ax=axes[1], edgecolor='black', color='#e67e22')
        axes[1].set_title('Log(Amount) Distribution', fontsize=14, fontweight='bold')
        axes[1].set_xlabel('Log(Amount)', fontsize=12)
        axes[1].set_ylabel('Frequency', fontsize=12)
        axes[1].grid(axis='y', alpha=0.3)
        
        plt.tight_layout()
        chart_paths['amount_dist'] = os.path.join(output_dir, 'amount_dist.png')
        plt.savefig(chart_paths['amount_dist'], dpi=300, bbox_inches='tight')
        plt.close()
    
    # 6. Feature Engineering Visualization (if processed data available)
    if fraud_processed is not None:
        print("📈 Creating feature engineering visualizations...")
        # Time-based features
        if 'hour_of_day' in fraud_processed.columns:
            fig, axes = plt.subplots(2, 2, figsize=(14, 10))
            
            # Hour distribution
            fraud_processed['hour_of_day'].hist(bins=24, ax=axes[0,0], edgecolor='black', color='#3498db')
            axes[0,0].set_title('Transaction Hour Distribution', fontsize=12, fontweight='bold')
            axes[0,0].set_xlabel('Hour of Day', fontsize=10)
            axes[0,0].set_ylabel('Frequency', fontsize=10)
            axes[0,0].grid(axis='y', alpha=0.3)
            
            # Day of week
            day_names = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
            if 'day_of_week' in fraud_processed.columns:
                day_counts = fraud_processed['day_of_week'].value_counts().sort_index()
                day_counts.index = [day_names[i] for i in day_counts.index]
                day_counts.plot(kind='bar', ax=axes[0,1], color='#2ecc71', edgecolor='black')
                axes[0,1].set_title('Transaction Day of Week', fontsize=12, fontweight='bold')
                axes[0,1].set_xlabel('Day', fontsize=10)
                axes[0,1].set_ylabel('Frequency', fontsize=10)
                axes[0,1].tick_params(axis='x', rotation=45)
                axes[0,1].grid(axis='y', alpha=0.3)
            
            # Weekend vs Weekday
            if 'is_weekend' in fraud_processed.columns:
                weekend_counts = fraud_processed['is_weekend'].value_counts()
                weekend_counts.index = ['Weekday', 'Weekend']
                weekend_counts.plot(kind='bar', ax=axes[1,0], color=['#f39c12', '#e74c3c'], edgecolor='black')
                axes[1,0].set_title('Weekend vs Weekday Transactions', fontsize=12, fontweight='bold')
                axes[1,0].set_xlabel('', fontsize=10)
                axes[1,0].set_ylabel('Count', fontsize=10)
                axes[1,0].tick_params(axis='x', rotation=0)
                axes[1,0].grid(axis='y', alpha=0.3)
            
            # Time since signup
            if 'time_since_signup' in fraud_processed.columns:
                fraud_processed['time_since_signup'].hist(bins=50, ax=axes[1,1], edgecolor='black', color='#9b59b6')
                axes[1,1].set_title('Time Since Signup (hours)', fontsize=12, fontweight='bold')
                axes[1,1].set_xlabel('Hours', fontsize=10)
                axes[1,1].set_ylabel('Frequency', fontsize=10)
                axes[1,1].set_xlim(0, 200)  # Focus on first 200 hours
                axes[1,1].grid(axis='y', alpha=0.3)
            
            plt.tight_layout()
            chart_paths['feature_engineering'] = os.path.join(output_dir, 'feature_engineering.png')
            plt.savefig(chart_paths['feature_engineering'], dpi=300, bbox_inches='tight')
            plt.close()
    
    # 7. Correlation Heatmap (Credit Card)
    print("📈 Creating correlation heatmap...")
    if 'Class' in creditcard_raw.columns:
        v_features = [col for col in creditcard_raw.columns if col.startswith('V')]
        if len(v_features) > 0:
            # Select top correlated features
            top_features = v_features[:10]  # First 10 V features
            corr_data = creditcard_raw[top_features + ['Class']].corr()
            
            plt.figure(figsize=(12, 10))
            sns.heatmap(corr_data, annot=True, fmt='.2f', cmap='coolwarm', center=0,
                       square=True, linewidths=0.5, cbar_kws={"shrink": 0.8})
            plt.title('Correlation Heatmap: Top V-Features vs Fraud (Credit Card)', 
                     fontsize=14, fontweight='bold', pad=20)
            plt.tight_layout()
            chart_paths['correlation_heatmap'] = os.path.join(output_dir, 'correlation_heatmap.png')
            plt.savefig(chart_paths['correlation_heatmap'], dpi=300, bbox_inches='tight')
            plt.close()
    
    # 8. Source and Browser Distribution
    if 'source' in fraud_raw.columns:
        fig, axes = plt.subplots(1, 2, figsize=(14, 5))
        
        fraud_raw['source'].value_counts().plot(kind='bar', ax=axes[0], color='#16a085', edgecolor='black')
        axes[0].set_title('Traffic Source Distribution', fontsize=14, fontweight='bold')
        axes[0].set_xlabel('Source', fontsize=12)
        axes[0].set_ylabel('Count', fontsize=12)
        axes[0].tick_params(axis='x', rotation=45)
        axes[0].grid(axis='y', alpha=0.3)
        
        if 'browser' in fraud_raw.columns:
            fraud_raw['browser'].value_counts().head(10).plot(kind='bar', ax=axes[1], 
                                                              color='#8e44ad', edgecolor='black')
            axes[1].set_title('Top 10 Browser Distribution', fontsize=14, fontweight='bold')
            axes[1].set_xlabel('Browser', fontsize=12)
            axes[1].set_ylabel('Count', fontsize=12)
            axes[1].tick_params(axis='x', rotation=45)
            axes[1].grid(axis='y', alpha=0.3)
        
        plt.tight_layout()
        chart_paths['source_browser'] = os.path.join(output_dir, 'source_browser.png')
        plt.savefig(chart_paths['source_browser'], dpi=300, bbox_inches='tight')
        plt.close()
    
    # 9. Class Distribution Before/After SMOTE
    print("📈 Creating SMOTE resampling comparison...")
    try:
        from sklearn.model_selection import train_test_split
        from imblearn.over_sampling import SMOTE
        
        if fraud_processed is not None and 'class' in fraud_processed.columns:
            # Prepare data for SMOTE demonstration
            X = fraud_processed.drop(columns=['class'])
            y = fraud_processed['class']
            
            # Get numeric columns only for SMOTE
            numeric_cols = X.select_dtypes(include=[np.number]).columns.tolist()
            X_numeric = X[numeric_cols].fillna(0)
            
            # Train-test split
            X_train, X_test, y_train, y_test = train_test_split(
                X_numeric, y, test_size=0.2, random_state=42, stratify=y
            )
            
            # Before SMOTE
            before_dist = y_train.value_counts().sort_index()
            
            # Apply SMOTE
            smote = SMOTE(random_state=42)
            X_train_smote, y_train_smote = smote.fit_resample(X_train, y_train)
            
            # After SMOTE
            after_dist = pd.Series(y_train_smote).value_counts().sort_index()
            
            # Create visualization
            fig, axes = plt.subplots(1, 2, figsize=(14, 6))
            
            # Before SMOTE
            before_dist.plot(kind='bar', ax=axes[0], color=['#2ecc71', '#e74c3c'], edgecolor='black')
            axes[0].set_title('Class Distribution BEFORE SMOTE (Training Set)', fontsize=14, fontweight='bold')
            axes[0].set_xlabel('Class (0=Legitimate, 1=Fraud)', fontsize=12)
            axes[0].set_ylabel('Count', fontsize=12)
            axes[0].tick_params(axis='x', rotation=0)
            axes[0].grid(axis='y', alpha=0.3)
            
            # Add count labels
            for i, (idx, val) in enumerate(before_dist.items()):
                axes[0].text(i, val, f'{val:,}', ha='center', va='bottom', fontweight='bold')
            
            imbalance_before = before_dist.max() / before_dist.min()
            axes[0].text(0.5, 0.95, f'Imbalance Ratio: {imbalance_before:.1f}:1', 
                        transform=axes[0].transAxes, ha='center', fontsize=11,
                        bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
            
            # After SMOTE
            after_dist.plot(kind='bar', ax=axes[1], color=['#2ecc71', '#e74c3c'], edgecolor='black')
            axes[1].set_title('Class Distribution AFTER SMOTE (Training Set)', fontsize=14, fontweight='bold')
            axes[1].set_xlabel('Class (0=Legitimate, 1=Fraud)', fontsize=12)
            axes[1].set_ylabel('Count', fontsize=12)
            axes[1].tick_params(axis='x', rotation=0)
            axes[1].grid(axis='y', alpha=0.3)
            
            # Add count labels
            for i, (idx, val) in enumerate(after_dist.items()):
                axes[1].text(i, val, f'{val:,}', ha='center', va='bottom', fontweight='bold')
            
            imbalance_after = 1.0  # Balanced after SMOTE
            axes[1].text(0.5, 0.95, f'Imbalance Ratio: {imbalance_after:.1f}:1 (Balanced)', 
                        transform=axes[1].transAxes, ha='center', fontsize=11,
                        bbox=dict(boxstyle='round', facecolor='lightgreen', alpha=0.5))
            
            plt.tight_layout()
            chart_paths['smote_comparison'] = os.path.join(output_dir, 'smote_comparison.png')
            plt.savefig(chart_paths['smote_comparison'], dpi=300, bbox_inches='tight')
            plt.close()
            
            # Credit card SMOTE comparison
            if 'Class' in creditcard_raw.columns:
                X_cc = creditcard_raw.drop(columns=['Class'])
                y_cc = creditcard_raw['Class']
                
                # Get numeric columns
                numeric_cols_cc = X_cc.select_dtypes(include=[np.number]).columns.tolist()
                X_cc_numeric = X_cc[numeric_cols_cc].fillna(0)
                
                # Train-test split
                X_train_cc, X_test_cc, y_train_cc, y_test_cc = train_test_split(
                    X_cc_numeric, y_cc, test_size=0.2, random_state=42, stratify=y_cc
                )
                
                # Before SMOTE
                before_dist_cc = y_train_cc.value_counts().sort_index()
                
                # Apply SMOTE
                X_train_cc_smote, y_train_cc_smote = smote.fit_resample(X_train_cc, y_train_cc)
                
                # After SMOTE
                after_dist_cc = pd.Series(y_train_cc_smote).value_counts().sort_index()
                
                # Create visualization
                fig, axes = plt.subplots(1, 2, figsize=(14, 6))
                
                # Before SMOTE
                before_dist_cc.plot(kind='bar', ax=axes[0], color=['#3498db', '#e74c3c'], edgecolor='black')
                axes[0].set_title('Credit Card - Class Distribution BEFORE SMOTE', fontsize=14, fontweight='bold')
                axes[0].set_xlabel('Class (0=Legitimate, 1=Fraud)', fontsize=12)
                axes[0].set_ylabel('Count', fontsize=12)
                axes[0].tick_params(axis='x', rotation=0)
                axes[0].grid(axis='y', alpha=0.3)
                
                # Add count labels
                for i, (idx, val) in enumerate(before_dist_cc.items()):
                    axes[0].text(i, val, f'{val:,}', ha='center', va='bottom', fontweight='bold', fontsize=9)
                
                imbalance_before_cc = before_dist_cc.max() / before_dist_cc.min()
                axes[0].text(0.5, 0.95, f'Imbalance Ratio: {imbalance_before_cc:.0f}:1', 
                            transform=axes[0].transAxes, ha='center', fontsize=11,
                            bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
                
                # After SMOTE
                after_dist_cc.plot(kind='bar', ax=axes[1], color=['#3498db', '#e74c3c'], edgecolor='black')
                axes[1].set_title('Credit Card - Class Distribution AFTER SMOTE', fontsize=14, fontweight='bold')
                axes[1].set_xlabel('Class (0=Legitimate, 1=Fraud)', fontsize=12)
                axes[1].set_ylabel('Count', fontsize=12)
                axes[1].tick_params(axis='x', rotation=0)
                axes[1].grid(axis='y', alpha=0.3)
                
                # Add count labels
                for i, (idx, val) in enumerate(after_dist_cc.items()):
                    axes[1].text(i, val, f'{val:,}', ha='center', va='bottom', fontweight='bold', fontsize=9)
                
                axes[1].text(0.5, 0.95, f'Imbalance Ratio: 1.0:1 (Balanced)', 
                            transform=axes[1].transAxes, ha='center', fontsize=11,
                            bbox=dict(boxstyle='round', facecolor='lightgreen', alpha=0.5))
                
                plt.tight_layout()
                chart_paths['smote_comparison_cc'] = os.path.join(output_dir, 'smote_comparison_cc.png')
                plt.savefig(chart_paths['smote_comparison_cc'], dpi=300, bbox_inches='tight')
                plt.close()
    except Exception as e:
        print(f"⚠ Could not create SMOTE comparison: {e}")
    
    # 10. Geolocation Fraud Pattern Analysis
    print("📈 Creating geolocation fraud pattern analysis...")
    try:
        if fraud_processed is not None and 'country' in fraud_processed.columns and 'class' in fraud_processed.columns:
            # Analyze fraud by country
            fraud_by_country = fraud_processed.groupby('country')['class'].agg(['mean', 'count']).sort_values('mean', ascending=False)
            fraud_by_country.columns = ['fraud_rate', 'transaction_count']
            fraud_by_country = fraud_by_country[fraud_by_country['transaction_count'] >= 10]  # At least 10 transactions
            
            # Create comprehensive visualization
            fig, axes = plt.subplots(2, 1, figsize=(14, 12))
            
            # Top 15 countries by fraud rate
            top_countries = fraud_by_country.head(15)
            axes[0].barh(range(len(top_countries)), top_countries['fraud_rate'], color='#e74c3c', alpha=0.7, edgecolor='black')
            axes[0].set_yticks(range(len(top_countries)))
            axes[0].set_yticklabels(top_countries.index)
            axes[0].set_xlabel('Fraud Rate', fontsize=12, fontweight='bold')
            axes[0].set_title('Top 15 Countries by Fraud Rate', fontsize=14, fontweight='bold', pad=20)
            axes[0].invert_yaxis()
            axes[0].grid(axis='x', alpha=0.3)
            axes[0].set_xlim(0, min(top_countries['fraud_rate'].max() * 1.15, 1.0))
            
            # Add fraud rate percentages as text
            for i, (country, row) in enumerate(top_countries.iterrows()):
                axes[0].text(row['fraud_rate'] + 0.005, i, f"{row['fraud_rate']*100:.1f}%", 
                            va='center', fontweight='bold', fontsize=10)
            
            # Top 15 countries by transaction volume with fraud rate
            top_volume = fraud_by_country.nlargest(15, 'transaction_count')
            colors_volume = ['#e74c3c' if rate > 0.1 else '#f39c12' if rate > 0.05 else '#3498db' 
                            for rate in top_volume['fraud_rate']]
            axes[1].barh(range(len(top_volume)), top_volume['fraud_rate'], color=colors_volume, alpha=0.7, edgecolor='black')
            axes[1].set_yticks(range(len(top_volume)))
            axes[1].set_yticklabels(top_volume.index)
            axes[1].set_xlabel('Fraud Rate', fontsize=12, fontweight='bold')
            axes[1].set_title('Top 15 Countries by Transaction Volume (with Fraud Rate)', 
                            fontsize=14, fontweight='bold', pad=20)
            axes[1].invert_yaxis()
            axes[1].grid(axis='x', alpha=0.3)
            axes[1].set_xlim(0, min(top_volume['fraud_rate'].max() * 1.15, 1.0))
            
            # Add transaction counts and fraud rates as text
            for i, (country, row) in enumerate(top_volume.iterrows()):
                axes[1].text(row['fraud_rate'] + 0.005, i, 
                            f"{row['fraud_rate']*100:.1f}% ({int(row['transaction_count']):,} txns)", 
                            va='center', fontsize=9, fontweight='bold')
            
            plt.tight_layout()
            chart_paths['geolocation_fraud'] = os.path.join(output_dir, 'geolocation_fraud.png')
            plt.savefig(chart_paths['geolocation_fraud'], dpi=300, bbox_inches='tight')
            plt.close()
            
            # Additional: Fraud rate distribution by country
            fig, ax = plt.subplots(figsize=(12, 6))
            fraud_by_country['fraud_rate'].hist(bins=30, ax=ax, edgecolor='black', color='#9b59b6', alpha=0.7)
            ax.set_title('Distribution of Fraud Rates Across Countries', fontsize=14, fontweight='bold')
            ax.set_xlabel('Fraud Rate', fontsize=12)
            ax.set_ylabel('Number of Countries', fontsize=12)
            ax.grid(axis='y', alpha=0.3)
            ax.axvline(fraud_by_country['fraud_rate'].mean(), color='red', linestyle='--', 
                      linewidth=2, label=f"Mean: {fraud_by_country['fraud_rate'].mean()*100:.2f}%")
            ax.axvline(fraud_by_country['fraud_rate'].median(), color='green', linestyle='--', 
                      linewidth=2, label=f"Median: {fraud_by_country['fraud_rate'].median()*100:.2f}%")
            ax.legend()
            plt.tight_layout()
            chart_paths['fraud_rate_distribution'] = os.path.join(output_dir, 'fraud_rate_distribution.png')
            plt.savefig(chart_paths['fraud_rate_distribution'], dpi=300, bbox_inches='tight')
            plt.close()
    except Exception as e:
        print(f"⚠ Could not create geolocation analysis: {e}")
    
    print(f"✓ Created {len(chart_paths)} visualizations")
    return chart_paths


def create_report_document(chart_paths, output_file='Fraud_Detection_Project_Report.docx'):
    """Create the DOCX report with all sections and visualizations"""
    
    base_path = os.path.dirname(os.path.dirname(__file__))
    data_raw = os.path.join(base_path, 'data', 'raw')
    data_processed = os.path.join(base_path, 'data', 'processed')
    
    # Load processed data for SMOTE comparison tables
    fraud_processed = None
    fraud_processed_path = os.path.join(data_processed, 'fraud_data_cleaned.csv')
    if os.path.exists(fraud_processed_path):
        try:
            fraud_processed = load_data(fraud_processed_path)
        except:
            pass
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('Fraud Detection for E-commerce and Bank Transactions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Comprehensive Project Report', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    company = doc.add_paragraph('Company: Adey Innovations Inc.')
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_format = company.runs[0].font
    company_format.bold = True
    
    date = doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Project Overview',
        '  2.1 Business Goals',
        '  2.2 Business Value and Customer Trust',
        '3. Datasets',
        '4. Data Loading & Inspection',
        '5. Data Cleaning',
        '6. Exploratory Data Analysis',
        '  6.4 Geolocation Fraud Pattern Analysis',
        '7. Feature Engineering',
        '8. Data Preprocessing',
        '9. Class Imbalance Handling',
        '  9.2 Impact of SMOTE Resampling',
        '10. Project Structure',
        '11. Key Findings & Insights',
        '12. Technical Implementation',
        '13. Next Steps',
        '  13.1 Task 2: Model Building and Evaluation',
        '  13.2 Task 3: Model Explainability',
        '  13.3 Anticipated Challenges and Mitigation Strategies',
        '  13.4 Additional Next Steps'
    ]
    
    for item in toc_items:
        if item.startswith('  '):
            # Sub-item with indentation
            p = doc.add_paragraph(item.strip(), style='List Bullet 2')
            p.style.font.size = Pt(10)
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            p = doc.add_paragraph(item, style='List Number')
            p.style.font.size = Pt(11)
    
    doc.add_page_break()
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This report documents the comprehensive work completed on the Fraud Detection project for '
        'E-commerce and Bank Transactions. The project focuses on building a robust, production-ready '
        'fraud detection system capable of identifying fraudulent transactions across two domains: '
        'e-commerce transactions and credit card payments.'
    )
    doc.add_paragraph(
        'Key accomplishments include:'
    )
    doc.add_paragraph('✓ Complete data cleaning and quality assurance', style='List Bullet')
    doc.add_paragraph('✓ Comprehensive exploratory data analysis with visualizations', style='List Bullet')
    doc.add_paragraph('✓ Advanced feature engineering (time-based, behavioral, statistical)', style='List Bullet')
    doc.add_paragraph('✓ Geolocation integration for fraud pattern analysis', style='List Bullet')
    doc.add_paragraph('✓ Class imbalance handling using SMOTE', style='List Bullet')
    doc.add_paragraph('✓ Production-grade modular code structure', style='List Bullet')
    doc.add_paragraph('✓ Complete preprocessing pipeline ready for model training', style='List Bullet')
    
    # Add summary statistics table
    doc.add_heading('1.1 Project Statistics Summary', level=2)
    try:
        fraud_raw = load_data(os.path.join(data_raw, 'Fraud_Data.csv'))
        creditcard_raw = load_data(os.path.join(data_raw, 'creditcard.csv'))
        
        stats_table = doc.add_table(rows=1, cols=4)
        stats_table.style = 'Light Grid Accent 1'
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = 'Dataset'
        hdr_cells[1].text = 'Total Transactions'
        hdr_cells[2].text = 'Features'
        hdr_cells[3].text = 'Fraud Rate'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Fraud data stats
        if 'class' in fraud_raw.columns:
            class_dist = fraud_raw['class'].value_counts()
            fraud_rate = (class_dist[1] / len(fraud_raw)) * 100 if 1 in class_dist.index else 0
            row_cells = stats_table.add_row().cells
            row_cells[0].text = 'E-commerce (Fraud_Data)'
            row_cells[1].text = f'{len(fraud_raw):,}'
            row_cells[2].text = f'{fraud_raw.shape[1]}'
            row_cells[3].text = f'{fraud_rate:.2f}%'
        
        # Credit card stats
        if 'Class' in creditcard_raw.columns:
            class_dist_cc = creditcard_raw['Class'].value_counts()
            fraud_rate_cc = (class_dist_cc[1] / len(creditcard_raw)) * 100 if 1 in class_dist_cc.index else 0
            row_cells = stats_table.add_row().cells
            row_cells[0].text = 'Credit Card'
            row_cells[1].text = f'{len(creditcard_raw):,}'
            row_cells[2].text = f'{creditcard_raw.shape[1]}'
            row_cells[3].text = f'{fraud_rate_cc:.2f}%'
    except:
        pass
    
    # 2. Project Overview
    doc.add_heading('2. Project Overview', level=1)
    doc.add_paragraph(
        'The primary objective of this project is to develop an advanced fraud detection system that can '
        'accurately identify fraudulent transactions while minimizing false positives. The system addresses '
        'two distinct transaction types:'
    )
    
    doc.add_paragraph('• E-commerce Transactions: User purchases with behavioral and geolocation data', style='List Bullet')
    doc.add_paragraph('• Credit Card Transactions: Anonymized bank transaction data with PCA features', style='List Bullet')
    
    doc.add_heading('2.1 Business Goals', level=2)
    goals_table = doc.add_table(rows=1, cols=2)
    goals_table.style = 'Light Grid Accent 1'
    hdr_cells = goals_table.rows[0].cells
    hdr_cells[0].text = 'Goal'
    hdr_cells[1].text = 'Description'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    
    goals = [
        ('High Precision & Recall', 'Accurately detect fraud while minimizing false positives'),
        ('Handle Class Imbalance', 'Effectively address severe class imbalance (10:1 to 580:1)'),
        ('Interpretable Results', 'Provide clear insights for business stakeholders'),
        ('Production-Ready Code', 'Build scalable, maintainable code structure'),
        ('Feature Engineering', 'Create meaningful features from raw transaction data')
    ]
    
    for goal, desc in goals:
        row_cells = goals_table.add_row().cells
        row_cells[0].text = goal
        row_cells[1].text = desc
    
    doc.add_heading('2.2 Business Value and Customer Trust', level=2)
    doc.add_paragraph(
        'Beyond preventing direct financial losses, accurate fraud detection plays a critical role in building '
        'and maintaining customer trust, which is essential for long-term business success. The impact extends '
        'far beyond immediate fraud prevention:'
    )
    
    trust_benefits = [
        ('Customer Confidence', 
         'When customers see that their transactions are protected by sophisticated fraud detection, they gain '
         'confidence in the platform\'s security. This trust encourages repeat purchases and increases customer '
         'lifetime value. Customers are more likely to recommend services they trust to others, leading to '
         'organic growth through word-of-mouth marketing.'),
        
        ('Reduced False Positives and Better User Experience',
         'By minimizing false positives (legitimate transactions flagged as fraud), we ensure that genuine '
         'customers experience seamless transactions without unnecessary friction. Every false positive creates '
         'frustration, potential customer service costs, and risks customer churn. A well-tuned fraud detection '
         'system balances security with convenience, showing customers that the platform understands and respects '
         'their legitimate behavior.'),
        
        ('Regulatory Compliance and Reputation',
         'Effective fraud detection helps organizations comply with financial regulations and industry standards. '
         'This compliance protects the company from legal issues and enhances its reputation in the marketplace. '
         'A strong reputation for security attracts new customers and business partners, creating additional '
         'revenue opportunities.'),
        
        ('Brand Protection',
         'Every fraudulent transaction that goes undetected can damage the brand\'s reputation when customers '
         'report unauthorized activity. Conversely, proactively preventing fraud demonstrates the company\'s '
         'commitment to customer protection, strengthening brand value and differentiation in competitive markets.'),
        
        ('Data-Driven Insights',
         'The fraud detection system provides valuable insights into transaction patterns, customer behavior, '
         'and emerging fraud trends. These insights can inform business strategy, product development, and '
         'marketing decisions, creating value beyond fraud prevention.')
    ]
    
    for benefit, description in trust_benefits:
        p = doc.add_paragraph()
        p.add_run(benefit + ': ').bold = True
        p.add_run(description)
    
    doc.add_paragraph(
        'The significance of class imbalance in fraud detection cannot be overstated. With fraud representing '
        'only 9% of e-commerce transactions and 0.17% of credit card transactions, naive models would achieve '
        'high accuracy by simply predicting "no fraud" for all cases. However, this approach fails to protect '
        'customers and the business. By addressing this imbalance through techniques like SMOTE and using '
        'appropriate evaluation metrics, we ensure that the system can reliably identify fraudulent patterns '
        'while maintaining customer trust through accurate, fair detection.'
    )
    
    # 3. Datasets
    doc.add_heading('3. Datasets', level=1)
    
    doc.add_heading('3.1 Fraud_Data.csv (E-commerce Transactions)', level=2)
    
    try:
        fraud_raw = load_data(os.path.join(data_raw, 'Fraud_Data.csv'))
        doc.add_paragraph(f'• Size: {fraud_raw.shape[0]:,} transactions × {fraud_raw.shape[1]} features')
        doc.add_paragraph('• Features: User information, transaction details, device & browser, IP address, traffic source')
        doc.add_paragraph('• Target: class (0 = Legitimate, 1 = Fraud)')
        
        if 'class' in fraud_raw.columns:
            class_dist = fraud_raw['class'].value_counts().sort_index()
            fraud_rate = (class_dist[1] / len(fraud_raw)) * 100 if 1 in class_dist.index else 0
            doc.add_paragraph(f'• Fraud Rate: {fraud_rate:.2f}% (Imbalance Ratio: {class_dist.max() / class_dist.min():.1f}:1)')
    except:
        doc.add_paragraph('• Size: ~151,000 transactions')
        doc.add_paragraph('• Fraud Rate: ~9% (Imbalance Ratio: ~10:1)')
    
    doc.add_heading('3.2 creditcard.csv (Bank Credit Card Transactions)', level=2)
    
    try:
        creditcard_raw = load_data(os.path.join(data_raw, 'creditcard.csv'))
        doc.add_paragraph(f'• Size: {creditcard_raw.shape[0]:,} transactions × {creditcard_raw.shape[1]} features')
        doc.add_paragraph('• Features: 28 anonymized V-features (PCA), transaction amount, time')
        doc.add_paragraph('• Target: Class (0 = Legitimate, 1 = Fraud)')
        
        if 'Class' in creditcard_raw.columns:
            class_dist_cc = creditcard_raw['Class'].value_counts().sort_index()
            fraud_rate_cc = (class_dist_cc[1] / len(creditcard_raw)) * 100 if 1 in class_dist_cc.index else 0
            doc.add_paragraph(f'• Fraud Rate: {fraud_rate_cc:.2f}% (Imbalance Ratio: {class_dist_cc.max() / class_dist_cc.min():.1f}:1)')
    except:
        doc.add_paragraph('• Size: ~284,000 transactions')
        doc.add_paragraph('• Fraud Rate: ~0.17% (Extreme Imbalance Ratio: ~580:1)')
    
    doc.add_heading('3.3 IpAddress_to_Country.csv', level=2)
    doc.add_paragraph('• Purpose: Map IP addresses to countries for geographical fraud analysis')
    doc.add_paragraph('• Features: IP ranges (lower_bound, upper_bound) and country codes')
    
    # Add class distribution charts
    if 'fraud_class_dist' in chart_paths and os.path.exists(chart_paths['fraud_class_dist']):
        doc.add_heading('3.4 Class Distribution Visualization', level=2)
        doc.add_paragraph('The following charts show the class imbalance in both datasets:')
        doc.add_picture(chart_paths['fraud_class_dist'], width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 1: E-commerce Transaction Class Distribution')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
        
        if 'creditcard_class_dist' in chart_paths and os.path.exists(chart_paths['creditcard_class_dist']):
            doc.add_picture(chart_paths['creditcard_class_dist'], width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('Figure 2: Credit Card Transaction Class Distribution')
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].runs[0].font.italic = True
    
    # 4. Data Loading & Inspection
    doc.add_heading('4. Data Loading & Inspection', level=1)
    doc.add_paragraph(
        'All data loading is handled through the reusable `data_loading.py` module, which provides:'
    )
    doc.add_paragraph('• File existence validation', style='List Bullet')
    doc.add_paragraph('• Comprehensive error handling', style='List Bullet')
    doc.add_paragraph('• Automatic shape and size reporting', style='List Bullet')
    doc.add_paragraph('• CSV parsing with error recovery', style='List Bullet')
    
    # 5. Data Cleaning
    doc.add_heading('5. Data Cleaning', level=1)
    doc.add_paragraph('Comprehensive data cleaning was performed using the `data_cleaning.py` module.')
    
    doc.add_heading('5.1 Cleaning Operations', level=2)
    cleaning_ops = [
        ('Duplicate Removal', 'Identified and removed duplicate records'),
        ('Missing Value Handling', 'Strategic imputation or dropping based on percentage (>50% missing = drop, else impute)'),
        ('Data Type Standardization', 'Converted columns to appropriate types (datetime, numeric)'),
        ('Invalid Data Filtering', 'Filtered invalid ages (0-120) and negative transaction values'),
        ('Critical Column Validation', 'Dropped rows with missing values in critical columns')
    ]
    
    cleaning_table = doc.add_table(rows=1, cols=2)
    cleaning_table.style = 'Light Grid Accent 1'
    hdr_cells = cleaning_table.rows[0].cells
    hdr_cells[0].text = 'Operation'
    hdr_cells[1].text = 'Description'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    
    for op, desc in cleaning_ops:
        row_cells = cleaning_table.add_row().cells
        row_cells[0].text = op
        row_cells[1].text = desc
    
    # 6. Exploratory Data Analysis
    doc.add_heading('6. Exploratory Data Analysis', level=1)
    doc.add_paragraph(
        'Comprehensive EDA was conducted to understand data distributions, relationships, and fraud patterns.'
    )
    
    doc.add_heading('6.1 Univariate Analysis', level=2)
    
    if 'purchase_value_dist' in chart_paths and os.path.exists(chart_paths['purchase_value_dist']):
        doc.add_picture(chart_paths['purchase_value_dist'], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 3: Purchase Value Distribution Analysis')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
    
    if 'age_dist' in chart_paths and os.path.exists(chart_paths['age_dist']):
        doc.add_picture(chart_paths['age_dist'], width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 4: Age Distribution')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
    
    if 'source_browser' in chart_paths and os.path.exists(chart_paths['source_browser']):
        doc.add_picture(chart_paths['source_browser'], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 5: Traffic Source and Browser Distribution')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
    
    doc.add_heading('6.2 Credit Card Transaction Analysis', level=2)
    
    if 'amount_dist' in chart_paths and os.path.exists(chart_paths['amount_dist']):
        doc.add_picture(chart_paths['amount_dist'], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 6: Transaction Amount Distribution (Original and Log-Transformed)')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
    
    doc.add_heading('6.3 Correlation Analysis', level=2)
    
    if 'correlation_heatmap' in chart_paths and os.path.exists(chart_paths['correlation_heatmap']):
        doc.add_picture(chart_paths['correlation_heatmap'], width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 7: Correlation Heatmap - Top V-Features vs Fraud')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
    
    doc.add_heading('6.4 Geolocation Fraud Pattern Analysis', level=2)
    doc.add_paragraph(
        'Geolocation analysis provides critical insights into fraud patterns by country, enabling targeted '
        'risk assessment and prevention strategies. IP address mapping to countries reveals geographic '
        'concentrations of fraudulent activity.'
    )
    
    if 'geolocation_fraud' in chart_paths and os.path.exists(chart_paths['geolocation_fraud']):
        doc.add_picture(chart_paths['geolocation_fraud'], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 9: Geolocation Fraud Pattern Analysis by Country')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
        
        doc.add_paragraph(
            'The top chart shows countries with the highest fraud rates, identifying high-risk regions that '
            'require enhanced monitoring. The bottom chart shows high-volume countries with their fraud rates, '
            'highlighting countries where even small fraud rate increases translate to significant financial '
            'impact due to transaction volume.'
        )
    
    if 'fraud_rate_distribution' in chart_paths and os.path.exists(chart_paths['fraud_rate_distribution']):
        doc.add_picture(chart_paths['fraud_rate_distribution'], width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 10: Distribution of Fraud Rates Across Countries')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
        
        doc.add_paragraph(
            'This histogram shows the overall distribution of fraud rates across all countries, with mean and '
            'median indicators. This helps understand the baseline fraud risk and identify outlier countries '
            'that deviate significantly from the norm.'
        )
    
    # Try to add specific country statistics if available
    try:
        fraud_raw = load_data(os.path.join(data_raw, 'Fraud_Data.csv'))
        ip_df = load_data(os.path.join(data_raw, 'IpAddress_to_Country.csv'))
        from feature_engineering import merge_geolocation
        
        if 'country' not in fraud_raw.columns:
            fraud_with_country = merge_geolocation(fraud_raw, ip_df)
        else:
            fraud_with_country = fraud_raw
        
        if 'country' in fraud_with_country.columns and 'class' in fraud_with_country.columns:
            fraud_by_country = fraud_with_country.groupby('country')['class'].agg(['mean', 'count']).sort_values('mean', ascending=False)
            fraud_by_country.columns = ['fraud_rate', 'transaction_count']
            fraud_by_country = fraud_by_country[fraud_by_country['transaction_count'] >= 10]
            
            doc.add_heading('6.4.1 Key Country-Level Insights', level=3)
            doc.add_paragraph('Top 5 Countries by Fraud Rate:')
            
            top_fraud_countries = fraud_by_country.head(5)
            country_table = doc.add_table(rows=1, cols=3)
            country_table.style = 'Light Grid Accent 1'
            hdr_cells = country_table.rows[0].cells
            hdr_cells[0].text = 'Country'
            hdr_cells[1].text = 'Fraud Rate'
            hdr_cells[2].text = 'Transaction Count'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            for country, row in top_fraud_countries.iterrows():
                row_cells = country_table.add_row().cells
                row_cells[0].text = str(country)
                row_cells[1].text = f"{row['fraud_rate']*100:.2f}%"
                row_cells[2].text = f"{int(row['transaction_count']):,}"
            
            doc.add_paragraph(
                'These insights enable the creation of country-based risk features that can enhance model '
                'performance by incorporating geographic fraud patterns into predictions.'
            )
    except:
        pass
    
    # 7. Feature Engineering
    doc.add_heading('7. Feature Engineering', level=1)
    doc.add_paragraph(
        'Advanced feature engineering was performed to create meaningful predictors for fraud detection.'
    )
    
    doc.add_heading('7.1 E-commerce Transaction Features', level=2)
    
    features_table = doc.add_table(rows=1, cols=3)
    features_table.style = 'Light Grid Accent 1'
    hdr_cells = features_table.rows[0].cells
    hdr_cells[0].text = 'Feature Category'
    hdr_cells[1].text = 'Features Created'
    hdr_cells[2].text = 'Rationale'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    feature_data = [
        ('Time-Based', 'hour_of_day, day_of_week, is_weekend, is_night', 'Fraudsters often operate during off-hours'),
        ('Behavioral', 'time_since_signup, quick_purchase, rapid_transactions', 'Quick actions and rapid transactions indicate suspicious behavior'),
        ('Transactional', 'transaction_count per user', 'Multiple transactions from same user can indicate fraud'),
        ('Geolocation', 'country (from IP mapping)', 'Certain countries show higher fraud rates')
    ]
    
    for category, features, rationale in feature_data:
        row_cells = features_table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = features
        row_cells[2].text = rationale
    
    doc.add_heading('7.2 Credit Card Transaction Features', level=2)
    
    cc_features_table = doc.add_table(rows=1, cols=3)
    cc_features_table.style = 'Light Grid Accent 1'
    hdr_cells = cc_features_table.rows[0].cells
    hdr_cells[0].text = 'Feature Category'
    hdr_cells[1].text = 'Features Created'
    hdr_cells[2].text = 'Rationale'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    cc_feature_data = [
        ('Statistical Aggregations', 'features_mean, features_std, features_min, features_max', 'Capture overall transaction behavior patterns from V-features'),
        ('Transformations', 'log_amount', 'Reduces skewness in transaction amounts'),
        ('Cyclical Encoding', 'hour_sin, hour_cos', 'Captures periodic patterns in transaction timing (24-hour cycle)')
    ]
    
    for category, features, rationale in cc_feature_data:
        row_cells = cc_features_table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = features
        row_cells[2].text = rationale
    
    if 'feature_engineering' in chart_paths and os.path.exists(chart_paths['feature_engineering']):
        doc.add_heading('7.3 Feature Engineering Visualizations', level=2)
        doc.add_picture(chart_paths['feature_engineering'], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 8: Time-Based Feature Engineering Visualizations')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
    
    # 8. Data Preprocessing
    doc.add_heading('8. Data Preprocessing', level=1)
    doc.add_paragraph('Preprocessing pipeline implemented using `preprocessing.py` module.')
    
    doc.add_heading('8.1 Preprocessing Steps', level=2)
    preprocessing_steps = [
        ('Train-Test Split', '80-20 split with stratification to maintain class distribution'),
        ('Feature Scaling', 'StandardScaler applied to numeric features (fit on training data only)'),
        ('Categorical Encoding', 'One-hot encoding ready for categorical variables'),
        ('Data Leakage Prevention', 'All transformations fitted on training data, then applied to test data')
    ]
    
    preproc_table = doc.add_table(rows=1, cols=2)
    preproc_table.style = 'Light Grid Accent 1'
    hdr_cells = preproc_table.rows[0].cells
    hdr_cells[0].text = 'Step'
    hdr_cells[1].text = 'Implementation'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    
    for step, impl in preprocessing_steps:
        row_cells = preproc_table.add_row().cells
        row_cells[0].text = step
        row_cells[1].text = impl
    
    # 9. Class Imbalance Handling
    doc.add_heading('9. Class Imbalance Handling', level=1)
    doc.add_paragraph(
        'Both datasets suffer from severe class imbalance, which would cause models to predict only the '
        'majority class if not addressed.'
    )
    
    doc.add_heading('9.1 SMOTE Implementation', level=2)
    doc.add_paragraph(
        'SMOTE (Synthetic Minority Over-sampling Technique) was applied to the training data only:'
    )
    doc.add_paragraph('• Creates synthetic samples in feature space (better than simple oversampling)', style='List Bullet')
    doc.add_paragraph('• Preserves original data distribution', style='List Bullet')
    doc.add_paragraph('• Applied ONLY to training data to prevent data leakage', style='List Bullet')
    doc.add_paragraph('• Test set maintains real-world imbalanced distribution for realistic evaluation', style='List Bullet')
    
    doc.add_heading('9.2 Impact of SMOTE Resampling', level=2)
    doc.add_paragraph(
        'The following visualizations demonstrate the significant impact of SMOTE resampling on class distribution. '
        'By comparing the before and after distributions, we can see how SMOTE effectively balances the classes, '
        'creating synthetic minority class samples in the feature space.'
    )
    
    if 'smote_comparison' in chart_paths and os.path.exists(chart_paths['smote_comparison']):
        doc.add_picture(chart_paths['smote_comparison'], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 11: Class Distribution Before and After SMOTE - E-commerce Data')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
        
        doc.add_paragraph(
            'The left panel shows the original imbalanced distribution in the training set, while the right panel '
            'shows the balanced distribution after SMOTE application. The imbalance ratio changes from approximately '
            '10:1 to 1:1, ensuring the model has sufficient examples of both classes for effective learning.'
        )
    
    if 'smote_comparison_cc' in chart_paths and os.path.exists(chart_paths['smote_comparison_cc']):
        doc.add_picture(chart_paths['smote_comparison_cc'], width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 12: Class Distribution Before and After SMOTE - Credit Card Data')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].runs[0].font.italic = True
        
        doc.add_paragraph(
            'For the credit card dataset, which has an extreme imbalance ratio of approximately 580:1, SMOTE '
            'dramatically increases the number of fraud examples in the training set. This extreme imbalance '
            'makes SMOTE particularly critical, as without it, models would achieve high accuracy by simply '
            'predicting "no fraud" for all transactions, completely failing to identify fraudulent cases.'
        )
    
    doc.add_heading('9.2.1 Quantitative Impact of SMOTE Resampling', level=3)
    doc.add_paragraph(
        'The following tables provide explicit quantitative documentation of SMOTE\'s impact on class distribution:'
    )
    
    # Create detailed comparison table for Fraud Data
    # Calculate SMOTE impact based on actual data distribution (no SMOTE execution needed)
    if fraud_processed is not None and 'class' in fraud_processed.columns:
        try:
            # Calculate expected values for 80-20 stratified split
            class_dist = fraud_processed['class'].value_counts().sort_index()
            total = len(fraud_processed)
            train_size = int(total * 0.8)
            
            # Stratified split maintains proportions
            legit_total = class_dist[0]
            fraud_total = class_dist[1]
            legit_train = int(legit_total * 0.8)
            fraud_train = int(fraud_total * 0.8)
            
            # After SMOTE: balance to 1:1 (majority class size)
            legit_after = legit_train  # No change
            fraud_after = legit_train  # Balanced to match majority
            total_after = legit_after + fraud_after
            
            before_dist = pd.Series({0: legit_train, 1: fraud_train})
            after_dist = pd.Series({0: legit_after, 1: fraud_after})
            
            doc.add_heading('E-commerce Dataset (Fraud_Data)', level=4)
            fraud_comparison_table = doc.add_table(rows=1, cols=5)
            fraud_comparison_table.style = 'Light Grid Accent 1'
            hdr_cells = fraud_comparison_table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Before SMOTE'
            hdr_cells[2].text = 'After SMOTE'
            hdr_cells[3].text = 'Change'
            hdr_cells[4].text = 'Impact'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Class 0 (Legitimate)
            row = fraud_comparison_table.add_row().cells
            row[0].text = 'Class 0 (Legitimate) Count'
            row[1].text = f"{before_dist[0]:,}"
            row[2].text = f"{after_dist[0]:,}"
            change = after_dist[0] - before_dist[0]
            row[3].text = f"{change:+,}"
            row[4].text = "No change (majority class)"
            
            # Class 1 (Fraud)
            row = fraud_comparison_table.add_row().cells
            row[0].text = 'Class 1 (Fraud) Count'
            row[1].text = f"{before_dist[1]:,}"
            row[2].text = f"{after_dist[1]:,}"
            change = after_dist[1] - before_dist[1]
            row[3].text = f"{change:+,} ({change/before_dist[1]*100:.1f}% increase)"
            row[4].text = "Synthetic samples created"
            
            # Total
            total_before = legit_train + fraud_train
            total_after = legit_after + fraud_after
            row = fraud_comparison_table.add_row().cells
            row[0].text = 'Total Samples'
            row[1].text = f"{total_before:,}"
            row[2].text = f"{total_after:,}"
            change = total_after - total_before
            row[3].text = f"{change:+,}"
            row[4].text = f"Dataset expanded by {change/total_before*100:.1f}%"
            
            # Imbalance Ratio
            ratio_before = before_dist.max() / before_dist.min()
            ratio_after = after_dist.max() / after_dist.min()
            row = fraud_comparison_table.add_row().cells
            row[0].text = 'Imbalance Ratio'
            row[1].text = f"{ratio_before:.2f}:1"
            row[2].text = f"{ratio_after:.2f}:1"
            change_pct = ((ratio_after - ratio_before) / ratio_before) * 100
            row[3].text = f"{change_pct:.1f}%"
            row[4].text = "Fully balanced (1:1)"
            
            # Fraud Percentage
            fraud_pct_before = (before_dist[1] / total_before) * 100
            fraud_pct_after = (after_dist[1] / total_after) * 100
            row = fraud_comparison_table.add_row().cells
            row[0].text = 'Fraud Percentage'
            row[1].text = f"{fraud_pct_before:.2f}%"
            row[2].text = f"{fraud_pct_after:.2f}%"
            change_pct_val = fraud_pct_after - fraud_pct_before
            row[3].text = f"{change_pct_val:+.2f}%"
            row[4].text = "Now 50% (balanced)"
        except Exception as e:
            doc.add_paragraph(f'Could not generate detailed SMOTE comparison table for fraud data: {str(e)}', style='Intense Quote')
    else:
        doc.add_paragraph('Fraud data not available for SMOTE comparison table generation.', style='Intense Quote')
    
    # Create detailed comparison table for Credit Card Data
    # Calculate SMOTE impact based on actual data distribution (no SMOTE execution needed)
    try:
        creditcard_raw = load_data(os.path.join(data_raw, 'creditcard.csv'))
        if 'Class' in creditcard_raw.columns:
            # Calculate expected values for 80-20 stratified split
            class_dist_cc = creditcard_raw['Class'].value_counts().sort_index()
            total_cc = len(creditcard_raw)
            
            legit_total_cc = class_dist_cc[0]
            fraud_total_cc = class_dist_cc[1]
            legit_train_cc = int(legit_total_cc * 0.8)
            fraud_train_cc = int(fraud_total_cc * 0.8)
            
            # After SMOTE: balance to 1:1 (majority class size)
            legit_after_cc = legit_train_cc  # No change
            fraud_after_cc = legit_train_cc  # Balanced to match majority
            total_after_cc = legit_after_cc + fraud_after_cc
            
            before_dist_cc = pd.Series({0: legit_train_cc, 1: fraud_train_cc})
            after_dist_cc = pd.Series({0: legit_after_cc, 1: fraud_after_cc})
            
            doc.add_heading('Credit Card Dataset', level=4)
            cc_comparison_table = doc.add_table(rows=1, cols=5)
            cc_comparison_table.style = 'Light Grid Accent 1'
            hdr_cells = cc_comparison_table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Before SMOTE'
            hdr_cells[2].text = 'After SMOTE'
            hdr_cells[3].text = 'Change'
            hdr_cells[4].text = 'Impact'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Class 0 (Legitimate)
            row = cc_comparison_table.add_row().cells
            row[0].text = 'Class 0 (Legitimate) Count'
            row[1].text = f"{before_dist_cc[0]:,}"
            row[2].text = f"{after_dist_cc[0]:,}"
            change = after_dist_cc[0] - before_dist_cc[0]
            row[3].text = f"{change:+,}"
            row[4].text = "No change (majority class)"
            
            # Class 1 (Fraud)
            row = cc_comparison_table.add_row().cells
            row[0].text = 'Class 1 (Fraud) Count'
            row[1].text = f"{before_dist_cc[1]:,}"
            row[2].text = f"{after_dist_cc[1]:,}"
            change = after_dist_cc[1] - before_dist_cc[1]
            row[3].text = f"{change:+,} ({change/before_dist_cc[1]*100:.1f}% increase)"
            row[4].text = "Synthetic samples created"
            
            # Total
            total_before_cc = legit_train_cc + fraud_train_cc
            total_after_cc = legit_after_cc + fraud_after_cc
            row = cc_comparison_table.add_row().cells
            row[0].text = 'Total Samples'
            row[1].text = f"{total_before_cc:,}"
            row[2].text = f"{total_after_cc:,}"
            change = total_after_cc - total_before_cc
            row[3].text = f"{change:+,}"
            row[4].text = f"Dataset expanded by {change/total_before_cc*100:.1f}%"
            
            # Imbalance Ratio
            ratio_before_cc = before_dist_cc.max() / before_dist_cc.min()
            ratio_after_cc = after_dist_cc.max() / after_dist_cc.min()
            row = cc_comparison_table.add_row().cells
            row[0].text = 'Imbalance Ratio'
            row[1].text = f"{ratio_before_cc:.1f}:1"
            row[2].text = f"{ratio_after_cc:.2f}:1"
            change_pct = ((ratio_after_cc - ratio_before_cc) / ratio_before_cc) * 100
            row[3].text = f"{change_pct:.1f}%"
            row[4].text = "Fully balanced (1:1)"
            
            # Fraud Percentage
            fraud_pct_before_cc = (before_dist_cc[1] / total_before_cc) * 100
            fraud_pct_after_cc = (after_dist_cc[1] / total_after_cc) * 100
            row = cc_comparison_table.add_row().cells
            row[0].text = 'Fraud Percentage'
            row[1].text = f"{fraud_pct_before_cc:.2f}%"
            row[2].text = f"{fraud_pct_after_cc:.2f}%"
            change_pct_val = fraud_pct_after_cc - fraud_pct_before_cc
            row[3].text = f"{change_pct_val:+.2f}%"
            row[4].text = "Now 50% (balanced)"
    except Exception as e:
        import traceback
        error_msg = f'Could not generate detailed SMOTE comparison table for credit card data: {str(e)}'
        doc.add_paragraph(error_msg, style='Intense Quote')
    
    doc.add_paragraph(
        'Key observations from the quantitative resampling results:'
    )
    doc.add_paragraph('• SMOTE successfully balances the class distribution to 1:1 ratio for training', style='List Bullet')
    doc.add_paragraph('• The synthetic samples are created in the feature space, not through simple duplication', style='List Bullet')
    doc.add_paragraph('• Test sets remain imbalanced to reflect real-world conditions for accurate evaluation', style='List Bullet')
    doc.add_paragraph('• The resampling is applied only to training data to prevent data leakage', style='List Bullet')
    doc.add_paragraph('• This balanced training set enables models to learn fraud patterns effectively', style='List Bullet')
    doc.add_paragraph(
        '• The dramatic increase in fraud samples (especially for credit card data with ~580:1 imbalance) '
        'demonstrates SMOTE\'s critical role in enabling effective fraud detection models', style='List Bullet'
    )
    
    doc.add_heading('9.3 Class Imbalance Statistics', level=2)
    
    imbalance_table = doc.add_table(rows=1, cols=4)
    imbalance_table.style = 'Light Grid Accent 1'
    hdr_cells = imbalance_table.rows[0].cells
    hdr_cells[0].text = 'Dataset'
    hdr_cells[1].text = 'Imbalance Ratio'
    hdr_cells[2].text = 'Fraud Rate'
    hdr_cells[3].text = 'Strategy'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    try:
        fraud_raw = load_data(os.path.join(data_raw, 'Fraud_Data.csv'))
        if 'class' in fraud_raw.columns:
            class_dist = fraud_raw['class'].value_counts().sort_index()
            fraud_rate = (class_dist[1] / len(fraud_raw)) * 100 if 1 in class_dist.index else 0
            ratio = class_dist.max() / class_dist.min()
            row_cells = imbalance_table.add_row().cells
            row_cells[0].text = 'E-commerce (Fraud_Data)'
            row_cells[1].text = f'{ratio:.1f}:1'
            row_cells[2].text = f'{fraud_rate:.2f}%'
            row_cells[3].text = 'SMOTE'
    except:
        row_cells = imbalance_table.add_row().cells
        row_cells[0].text = 'E-commerce (Fraud_Data)'
        row_cells[1].text = '~10:1'
        row_cells[2].text = '~9%'
        row_cells[3].text = 'SMOTE'
    
    try:
        creditcard_raw = load_data(os.path.join(data_raw, 'creditcard.csv'))
        if 'Class' in creditcard_raw.columns:
            class_dist_cc = creditcard_raw['Class'].value_counts().sort_index()
            fraud_rate_cc = (class_dist_cc[1] / len(creditcard_raw)) * 100 if 1 in class_dist_cc.index else 0
            ratio_cc = class_dist_cc.max() / class_dist_cc.min()
            row_cells = imbalance_table.add_row().cells
            row_cells[0].text = 'Credit Card'
            row_cells[1].text = f'{ratio_cc:.1f}:1'
            row_cells[2].text = f'{fraud_rate_cc:.2f}%'
            row_cells[3].text = 'SMOTE'
    except:
        row_cells = imbalance_table.add_row().cells
        row_cells[0].text = 'Credit Card'
        row_cells[1].text = '~580:1'
        row_cells[2].text = '~0.17%'
        row_cells[3].text = 'SMOTE'
    
    # 10. Project Structure
    doc.add_heading('10. Project Structure', level=1)
    doc.add_paragraph('The project follows a modular, production-grade structure:')
    
    structure_text = """src/
├── data_loading.py          # Data loading with error handling
├── data_cleaning.py         # Data cleaning and standardization
├── feature_engineering.py   # Feature engineering and geolocation
├── eda_utils.py             # EDA utilities and visualizations
└── preprocessing.py         # Preprocessing pipeline (scaling, encoding, SMOTE)

notebooks/
└── eda-fraud-data.ipynb     # Comprehensive EDA and preprocessing notebook

data/
├── raw/                     # Original datasets (read-only)
│   ├── Fraud_Data.csv
│   ├── creditcard.csv
│   └── IpAddress_to_Country.csv
└── processed/               # Cleaned and preprocessed data
    ├── fraud_data_cleaned.csv
    ├── creditcard_cleaned.csv
    └── (train/test splits with SMOTE)"""
    
    p = doc.add_paragraph(structure_text, style='No Spacing')
    p.style.font.name = 'Courier New'
    p.style.font.size = Pt(9)
    
    # 11. Key Findings & Insights
    doc.add_heading('11. Key Findings & Insights', level=1)
    
    doc.add_heading('11.1 E-commerce Transactions (Fraud_Data)', level=2)
    findings_fraud = [
        'Imbalance Ratio: ~10:1 (legitimate:fraud)',
        'High-risk patterns identified:',
        '  • Very quick purchases after signup (<5 minutes)',
        '  • Late night transactions (10 PM - 6 AM)',
        '  • Weekend transactions',
        '  • Multiple rapid transactions from same user',
        '  • Certain countries show higher fraud rates'
    ]
    
    for finding in findings_fraud:
        if finding.startswith('  •'):
            doc.add_paragraph(finding, style='List Bullet 2')
        elif finding.endswith(':'):
            doc.add_paragraph(finding, style='List Bullet')
        else:
            doc.add_paragraph(finding, style='List Bullet')
    
    doc.add_heading('11.2 Credit Card Transactions', level=2)
    findings_cc = [
        'Imbalance Ratio: ~580:1 (extreme imbalance!)',
        'High-risk patterns identified:',
        '  • Unusual transaction amounts',
        '  • Atypical V-feature values (PCA components)',
        '  • Transaction timing patterns',
        '  • Statistical aggregations capture overall behavior'
    ]
    
    for finding in findings_cc:
        if finding.startswith('  •'):
            doc.add_paragraph(finding, style='List Bullet 2')
        elif finding.endswith(':'):
            doc.add_paragraph(finding, style='List Bullet')
        else:
            doc.add_paragraph(finding, style='List Bullet')
    
    # 12. Technical Implementation
    doc.add_heading('12. Technical Implementation', level=1)
    
    doc.add_heading('12.1 Technologies Used', level=2)
    tech_table = doc.add_table(rows=1, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text = 'Technology'
    hdr_cells[1].text = 'Purpose'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    
    technologies = [
        ('pandas', 'Data manipulation and analysis'),
        ('numpy', 'Numerical computing'),
        ('scikit-learn', 'Machine learning algorithms and preprocessing'),
        ('imbalanced-learn', 'SMOTE and class imbalance handling'),
        ('matplotlib & seaborn', 'Data visualization'),
        ('Jupyter Notebook', 'Interactive exploration and analysis')
    ]
    
    for tech, purpose in technologies:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose
    
    doc.add_heading('12.2 Best Practices Implemented', level=2)
    best_practices = [
        'Modular Design: Reusable functions in src/ modules',
        'Error Handling: Comprehensive try/except blocks with informative messages',
        'Data Leakage Prevention: Train-test split before any transformation',
        'Documentation: Clear docstrings and inline comments',
        'Reproducibility: Random seeds set for all random operations',
        'Code Quality: Functions over repeated code, clear naming conventions'
    ]
    
    for practice in best_practices:
        doc.add_paragraph(practice, style='List Bullet')
    
    # 13. Next Steps
    doc.add_heading('13. Next Steps', level=1)
    doc.add_paragraph('The following steps are recommended to complete the fraud detection system:')
    
    doc.add_heading('13.1 Task 2: Model Building and Evaluation', level=2)
    
    doc.add_paragraph('Model Training:', style='Heading 3')
    doc.add_paragraph(
        'Use SMOTE-balanced training sets to train multiple models including Logistic Regression (baseline), '
        'Random Forest, XGBoost, and Neural Networks. Each model offers different strengths: Logistic Regression '
        'provides interpretability, Random Forest handles non-linear relationships, XGBoost offers superior '
        'performance with gradient boosting, and Neural Networks can capture complex patterns.'
    )
    
    doc.add_paragraph('Model Evaluation:', style='Heading 3')
    doc.add_paragraph(
        'Evaluate models on imbalanced test sets (real-world distribution) using metrics appropriate for '
        'imbalanced data: Precision (minimize false positives), Recall (capture all fraud cases), F1-Score '
        '(balance of precision and recall), ROC-AUC (overall discrimination ability), and PR-AUC (especially '
        'important for imbalanced data as it focuses on the minority class).'
    )
    
    doc.add_heading('13.2 Task 3: Model Explainability', level=2)
    doc.add_paragraph(
        'Implement SHAP (SHapley Additive exPlanations) values to understand model decisions and provide '
        'explanations for fraud predictions. This is critical for regulatory compliance, customer service '
        '(explaining why transactions were flagged), and model debugging. Additional explainability techniques '
        'include feature importance analysis, LIME (Local Interpretable Model-agnostic Explanations), and '
        'decision tree visualization for tree-based models.'
    )
    
    doc.add_heading('13.3 Anticipated Challenges and Mitigation Strategies', level=2)
    
    doc.add_paragraph(
        'The following table outlines key challenges anticipated in the next phases and proposed mitigation strategies:'
    )
    
    challenges_table = doc.add_table(rows=1, cols=3)
    challenges_table.style = 'Light Grid Accent 1'
    hdr_cells = challenges_table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Key Challenges'
    hdr_cells[2].text = 'Mitigation Strategies'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    challenges_data = [
        ('Model Building\n(Task 2)',
         '• Extreme class imbalance may still cause models to favor majority class despite SMOTE\n'
         '• Overfitting on SMOTE-synthetic samples\n'
         '• Model selection trade-offs (accuracy vs interpretability vs computational cost)\n'
         '• Feature space curse of dimensionality with many engineered features',
         '• Use stratified cross-validation with SMOTE applied only to training folds\n'
         '• Combine SMOTE with class weights or ensemble methods\n'
         '• Implement early stopping, regularization, and cross-validation to prevent overfitting\n'
         '• Perform feature selection using correlation analysis and recursive feature elimination\n'
         '• Compare models using multiple metrics, not just accuracy\n'
         '• Consider cost-sensitive learning approaches'),
        
        ('Model Evaluation\n(Task 2)',
         '• Test set performance may differ significantly from training performance due to imbalance\n'
         '• Need to balance precision and recall based on business requirements\n'
         '• Evaluation metrics may be misleading if not chosen carefully',
         '• Evaluate on imbalanced test set (real-world scenario) not SMOTE-balanced data\n'
         '• Use confusion matrix to analyze false positives vs false negatives\n'
         '• Establish business-specific thresholds (e.g., prioritize recall for fraud detection)\n'
         '• Create cost-benefit analysis: cost of false positive vs cost of missed fraud\n'
         '• Use PR-AUC instead of ROC-AUC for highly imbalanced data'),
        
        ('Model Explainability\n(Task 3)',
         '• Complex models (XGBoost, Neural Networks) are less interpretable\n'
         '• SHAP computation can be computationally expensive for large datasets\n'
         '• Explaining predictions to non-technical stakeholders\n'
         '• Balancing model complexity with explainability requirements',
         '• Use SHAP TreeExplainer for tree-based models (faster computation)\n'
         '• Sample representative cases for explanation to reduce computational load\n'
         '• Create visualizations and summary statistics (e.g., average SHAP values by feature)\n'
         '• Develop clear, non-technical explanations using feature importance rankings\n'
         '• Consider ensemble of interpretable models if explainability is critical\n'
         '• Document common fraud patterns identified by SHAP for business understanding'),
        
        ('Production Deployment',
         '• Real-time prediction latency requirements\n'
         '• Model drift over time as fraud patterns evolve\n'
         '• Handling new transaction types or features not seen in training\n'
         '• Scaling to handle high transaction volumes',
         '• Optimize models for inference speed (feature pruning, model compression)\n'
         '• Implement model versioning and A/B testing framework\n'
         '• Set up continuous monitoring with performance dashboards\n'
         '• Schedule periodic model retraining (e.g., monthly or quarterly)\n'
         '• Implement feature validation pipeline to catch data quality issues\n'
         '• Use distributed computing or model serving frameworks (e.g., TensorFlow Serving, MLflow)'),
        
        ('Data Quality\n& Maintenance',
         '• Missing or corrupted real-time transaction data\n'
         '• IP geolocation mapping may become outdated\n'
         '• Changes in user behavior patterns over time',
         '• Implement robust data validation and error handling\n'
         '• Maintain updated geolocation mapping databases\n'
         '• Monitor feature distributions for drift detection\n'
         '• Establish data quality metrics and alerting systems')
    ]
    
    for phase, challenges, mitigations in challenges_data:
        row_cells = challenges_table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = challenges
        row_cells[2].text = mitigations
    
    doc.add_heading('13.4 Additional Next Steps', level=2)
    additional_steps = [
        'Feature Selection: Perform correlation analysis and feature importance analysis to select optimal features and reduce dimensionality',
        'Hyperparameter Tuning: Use grid search, random search, or Bayesian optimization to optimize model parameters',
        'Production Deployment: Create API endpoints using FastAPI for real-time fraud detection with proper authentication and rate limiting',
        'Monitoring: Set up monitoring and logging for model performance in production, including prediction latency, accuracy metrics, and data drift detection',
        'Documentation: Create comprehensive model documentation including training procedures, feature descriptions, and deployment guides',
        'A/B Testing Framework: Implement framework to compare model versions and measure business impact'
    ]
    
    for step in additional_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    # Save document
    output_path = os.path.join(base_path, output_file)
    doc.save(output_path)
    print(f"\n✓ Report saved to: {output_path}")
    return output_path


def main():
    """Main function to generate the report"""
    print("="*60)
    print("Generating Comprehensive Project Report")
    print("="*60)
    
    # Create visualizations
    chart_paths = create_visualizations()
    
    # Create document
    output_file = create_report_document(chart_paths)
    
    print("\n" + "="*60)
    print(f"✓ Report generation complete!")
    print(f"✓ Output file: {output_file}")
    print("="*60)


if __name__ == '__main__':
    main()

